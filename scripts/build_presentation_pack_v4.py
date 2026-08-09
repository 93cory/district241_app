from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
OUT.mkdir(exist_ok=True)

GABON_GREEN = RGBColor(0, 98, 51)
GABON_BLUE = RGBColor(12, 126, 180)
GOLD = RGBColor(248, 197, 55)


def set_run(run, *, size: int | None = None, bold: bool = False, color: RGBColor | None = None):
    run.font.name = "Aptos"
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def title(doc: Document, text: str, subtitle: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("REPUBLIQUE GABONAISE\n")
    set_run(r, size=11, bold=True, color=GABON_GREEN)
    r = p.add_run("Ministere de l'Industrie\n")
    set_run(r, size=13, bold=True, color=GABON_BLUE)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run(r, size=22, bold=True, color=GABON_GREEN)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run(r, size=12, color=RGBColor(82, 97, 117))
    doc.add_paragraph()


def h(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, size=15, bold=True, color=GABON_GREEN)


def p(doc: Document, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.08
    run = para.add_run(text)
    set_run(run, size=10)


def bullets(doc: Document, items: list[str]):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(item)
        set_run(run, size=10)


def table(doc: Document, headers: list[str], rows: list[list[str]]):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        run = t.rows[0].cells[i].paragraphs[0].add_run(header)
        set_run(run, size=9, bold=True, color=GABON_GREEN)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            run = cells[i].paragraphs[0].add_run(value)
            set_run(run, size=9)
    doc.add_paragraph()


def footer(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    f = section.footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = f.add_run("PNPI - Document de travail pour presentation ministerielle")
    set_run(r, size=8, color=RGBColor(82, 97, 117))


def make_note() -> Path:
    doc = Document()
    footer(doc)
    title(
        doc,
        "Note conceptuelle PNPI - Version 4",
        "Plateforme Nationale de Pilotage Industriel avec extensions AGANOR et OGAPI",
    )

    h(doc, "1. Objet")
    p(
        doc,
        "La PNPI vise a doter le Ministere de l'Industrie d'un outil numerique unique pour piloter "
        "les agrements techniques industriels, les inspections, les indicateurs de performance et la "
        "coordination progressive avec les institutions sectorielles rattachees ou partenaires.",
    )

    h(doc, "2. Vision institutionnelle")
    bullets(
        doc,
        [
            "Un cockpit ministeriel pour suivre les dossiers, les delais, les alertes et les arbitrages.",
            "Un guichet operateur pour deposer, completer et suivre les demandes d'agrement.",
            "Un espace AGANOR pour projeter la normalisation, la certification et la conformite produit.",
            "Un espace OGAPI pour projeter la propriete industrielle, les marques, brevets et dessins industriels.",
            "Une gouvernance des donnees permettant de produire des rapports fiables pour le Ministre et les directions.",
        ],
    )

    h(doc, "3. Perimetre fonctionnel prioritaire")
    table(
        doc,
        ["Volet", "Fonctionnalites", "Impact attendu"],
        [
            [
                "ATI",
                "Depot, instruction, pieces, workflow, decision, historique, QR code.",
                "Reduction des delais et tracabilite des decisions.",
            ],
            [
                "Inspections",
                "Planification, rapports terrain, photos, geolocalisation, suivi des reserves.",
                "Controle plus objectif et preuves terrain centralisees.",
            ],
            [
                "Pilotage",
                "KPI, carte nationale, alertes SLA, rapports, synthese executive.",
                "Lecture rapide de l'etat industriel national.",
            ],
            [
                "AGANOR",
                "Normalisation, certifications, laboratoires, standards applicables.",
                "Lien entre agrement industriel et exigences qualite.",
            ],
            [
                "OGAPI",
                "Propriete industrielle, portefeuille de titres, alertes d'echeance.",
                "Protection et valorisation des actifs immateriels.",
            ],
        ],
    )

    h(doc, "4. Proposition pour la presentation")
    bullets(
        doc,
        [
            "Ouvrir par le besoin ministeriel : voir, decider, tracer.",
            "Montrer l'application deja fonctionnelle : connexion, cockpit, dossier ATI, inspection, institutions.",
            "Positionner AGANOR et OGAPI comme extensions structurees et prudentes, sans pretendre remplacer leurs systemes existants.",
            "Conclure sur une feuille de route en lots : stabilisation demo, cadrage institutionnel, pilote, deploiement progressif.",
        ],
    )

    h(doc, "5. Decision attendue")
    p(
        doc,
        "Valider le principe d'une demonstration ministerielle, designer les points focaux techniques et "
        "autoriser une phase courte de cadrage avec les delegues concernes afin de confirmer les donnees, "
        "les responsabilites et les interfaces prioritaires.",
    )

    path = OUT / "Note_conceptuelle_PNPI_V4_AGANOR_OGAPI.docx"
    doc.save(path)
    return path


def make_conducteur() -> Path:
    doc = Document()
    footer(doc)
    title(
        doc,
        "Conducteur de demonstration Ministre",
        "Parcours court, messages cles et comptes de demonstration",
    )

    h(doc, "1. Objectif de la demo")
    p(
        doc,
        "Demontrer que la PNPI peut devenir le tableau de bord operationnel du Ministere : une plateforme "
        "qui transforme les dossiers industriels en decisions suivies, mesurables et traçables.",
    )

    h(doc, "2. Comptes de demonstration")
    table(
        doc,
        ["Profil", "Identifiant", "Mot de passe"],
        [
            ["Administrateur", "admin", "Demo1234!@#$"],
            ["Ministre", "ministre", "Demo1234!@#$"],
            ["Directeur", "directeur", "Demo1234!@#$"],
            ["Instructeur", "instructeur", "Demo1234!@#$"],
            ["Inspecteur", "inspecteur", "Demo1234!@#$"],
            ["Operateur", "operateur", "Demo1234!@#$"],
        ],
    )

    h(doc, "3. Parcours conseille en 12 minutes")
    table(
        doc,
        ["Temps", "Ecran", "Message a porter"],
        [
            ["1 min", "/", "La PNPI est la porte d'entree numerique du pilotage industriel."],
            ["1 min", "/connexion", "Connexion securisee par profil : Ministre, directions, instructeurs, inspecteurs, operateurs."],
            ["3 min", "/pnpi", "Vue Ministre : KPI, alertes, carte, dossiers recents, et priorites du jour."],
            ["2 min", "/pnpi/executive", "Synthese executive pour arbitrage et reporting institutionnel."],
            ["2 min", "/pnpi/ati", "Instruction ATI : dossier, workflow, pieces, historique et decision."],
            ["1 min", "/pnpi/inspections", "Controle terrain : planification, rapport, photos, preuves."],
            ["2 min", "/pnpi/institutions/cockpit", "Extension institutionnelle : Ministere, AGANOR, OGAPI, sans surpromesse technique."],
        ],
    )

    h(doc, "4. Points de vigilance pendant la presentation")
    bullets(
        doc,
        [
            "Eviter de presenter AGANOR ou OGAPI comme des modules finalises : parler de trajectoire d'integration.",
            "Insister sur la traçabilite, les delais, la transparence et l'aide a la decision.",
            "Garder un parcours court : le Ministre doit comprendre la valeur avant les details techniques.",
            "Prevoir une solution de secours : captures d'ecran ou maquettes locales si le reseau est instable.",
        ],
    )

    h(doc, "5. Questions probables")
    table(
        doc,
        ["Question", "Reponse courte proposee"],
        [
            [
                "Est-ce deja operationnel ?",
                "Un socle fonctionnel existe pour la demo ; le deploiement officiel doit passer par cadrage, securite et validation metier.",
            ],
            [
                "Que devient AGANOR ?",
                "AGANOR conserve son metier ; la PNPI peut exposer une vue de coordination normes/certifications.",
            ],
            [
                "Que devient OGAPI ?",
                "OGAPI conserve son metier ; la PNPI peut relier les projets industriels aux titres de propriete industrielle pertinents.",
            ],
            [
                "Quel est le prochain pas ?",
                "Nommer les points focaux, valider les donnees prioritaires et organiser un pilote limite.",
            ],
        ],
    )

    path = OUT / "Conducteur_demo_Ministre_PNPI_AGANOR_OGAPI.docx"
    doc.save(path)
    return path


def export_pdf(docx_path: Path) -> Path | None:
    try:
        import pythoncom
        import win32com.client
    except Exception:
        return None

    pdf_path = docx_path.with_suffix(".pdf")
    pythoncom.CoInitialize()
    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path))
        doc.SaveAs(str(pdf_path), FileFormat=17)
        doc.Close(False)
        return pdf_path
    finally:
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def main():
    created = [make_note(), make_conducteur()]
    for path in created:
        print(path)
        pdf = export_pdf(path)
        if pdf:
            print(pdf)
        else:
            print(f"PDF export skipped for {path.name}: Word COM unavailable")


if __name__ == "__main__":
    main()
