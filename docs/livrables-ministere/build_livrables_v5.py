from __future__ import annotations

import re
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = Path("D:/pnpi/outputs/livrables-ministere-v5")
OUTPUT_SUFFIX = os.getenv("PNPI_LIVRABLES_SUFFIX", "V5")
OUTPUT_DOCX = OUTPUT_DIR / f"PNPI_Livrables_Ministere_{OUTPUT_SUFFIX}.docx"

SOURCE_FILES = [
    "00-plan-maitre-livrables.md",
    "01-dossier-executif-v5.md",
    "02-presentation-executive-v5.md",
    "03-conducteur-demo-v5.md",
    "04-architecture-technique-cible-v5.md",
    "05-plan-developpement-budget-ressources-v5.md",
    "06-infrastructure-cybersecurite-deploiement-v5.md",
    "07-cahier-des-charges-v5.md",
    "08-annexes-domaines-metiers-v5.md",
    "09-budget-reel-strategie-contractualisation-v5.md",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(clean_inline(text))
    run.bold = bold
    run.font.size = Pt(9)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def clean_inline(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("`", "")
    return text.strip()


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2F3")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_code_block(doc: Document, lines: list[str]) -> None:
    if not lines:
        return
    p = doc.add_paragraph()
    p.style = "CodeBlock"
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(31, 78, 121)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx].cells[c_idx]
            set_cell_text(cell, text, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
    doc.add_paragraph()


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and is_table_line(lines[i]):
        raw_cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in raw_cells):
            rows.append(raw_cells)
        i += 1
    return rows, i


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Heading 1", 16, "003F8F"),
        ("Heading 2", 13, "003F8F"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    code = doc.styles.add_style("CodeBlock", 1)
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code.paragraph_format.left_indent = Cm(0.25)
    code.paragraph_format.right_indent = Cm(0.25)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run("Plateforme Nationale de Pilotage Industriel")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 63, 143)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PNPI — Dossier consolidé des livrables ministériels V5")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0, 148, 64)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("Ministère de l’Industrie — République Gabonaise")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    run = p.add_run("Version de travail — 29 juillet 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(99, 99, 99)

    add_horizontal_rule(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Ce dossier regroupe le plan maître, le dossier exécutif, la trame de présentation, "
        "le conducteur de démonstration, l’architecture cible, le plan de développement, "
        "l’infrastructure/cybersécurité, le cahier des charges et les annexes domaines métiers."
    )
    run.font.size = Pt(11)
    run.font.italic = True
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("Sommaire", level=1)
    for idx, filename in enumerate(SOURCE_FILES, start=1):
        title = (ROOT / filename).read_text(encoding="utf-8").splitlines()[0].lstrip("# ").strip()
        p = doc.add_paragraph(style="List Number")
        p.add_run(title).bold = True
    doc.add_page_break()


def add_markdown_file(doc: Document, path: Path, first: bool = False) -> None:
    if not first:
        doc.add_section(WD_SECTION_START.NEW_PAGE)
    lines = path.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped == "---":
            add_horizontal_rule(doc)
            i += 1
            continue
        if is_table_line(line):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if stripped.startswith("### "):
            doc.add_heading(clean_inline(stripped[4:]), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(clean_inline(stripped[3:]), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(clean_inline(stripped[2:]), level=1)
        elif re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(clean_inline(re.sub(r"^[-*]\s+", "", stripped)))
        elif re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            p.add_run(clean_inline(re.sub(r"^\d+\.\s+", "", stripped)))
        elif stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(clean_inline(stripped.lstrip("> ")))
            run.italic = True
            run.font.color.rgb = RGBColor(31, 78, 121)
        else:
            p = doc.add_paragraph()
            p.add_run(clean_inline(stripped))
        i += 1
    if code_lines:
        add_code_block(doc, code_lines)


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("PNPI — Livrables ministériels V5")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(99, 99, 99)


def build() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_toc(doc)
    for index, filename in enumerate(SOURCE_FILES):
        add_markdown_file(doc, ROOT / filename, first=(index == 0))
    add_footer(doc)
    doc.core_properties.title = "PNPI — Livrables ministériels V5"
    doc.core_properties.subject = "Dossier consolidé des livrables PNPI"
    doc.core_properties.author = "PNPI / Codex"
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build()
