from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = Path("D:/pnpi/outputs/livrables-ministere-v5")
SOURCE = ROOT / "10-note-protection-cadrage-pnpi-v5.md"
OUTPUT_DOCX = OUTPUT_DIR / "PNPI_Note_Protection_Cadrage_V5.docx"


def clean(text: str) -> str:
    return (
        text.replace("—", "-")
        .replace("’", "'")
        .replace("“", '"')
        .replace("”", '"')
        .strip()
    )


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(clean(text))
    run.bold = bold
    run.font.size = Pt(9)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyper_simple_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(clean(text))
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    return p


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(clean(text))
    run.font.name = "Arial"
    run.font.size = Pt(9.5)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(clean(text))
    run.font.name = "Arial"
    run.font.size = Pt(9.5)


def parse_inline(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(clean(part[2:-2]))
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(clean(part[1:-1]))
            run.bold = True
            run.font.name = "Consolas"
        else:
            run = paragraph.add_run(clean(part))
        run.font.name = run.font.name or "Arial"
        run.font.size = Pt(9.5)


def add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(clean(text))
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(31, 77, 120)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(2.1), Inches(4.2)]
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.width = widths[col_idx]
            set_cell_text(cell, value, bold=row_idx == 0)
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
    doc.add_paragraph()


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 12.5, RGBColor(46, 116, 181)),
        ("Heading 2", 11, RGBColor(46, 116, 181)),
        ("Heading 3", 10.5, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(5)
        style.paragraph_format.space_after = Pt(2)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "PNPI - Note de protection du portage initial"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    footer = section.footer.paragraphs[0]
    footer.text = "Document de travail - a joindre aux pieces de preuve et au dossier executif"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.name = "Arial"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(100, 116, 139)


def build() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("PNPI - Note de protection du portage initial et proposition de cadrage")
    run.font.name = "Arial"
    run.font.size = Pt(15.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("M. MBA NDONG Jean Baptiste - Version V5 - 29 juillet 2026")
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(100, 116, 139)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_table = False
    table_rows: list[list[str]] = []
    skip_title_block = True

    for raw in lines:
        line = raw.rstrip()
        if skip_title_block:
            if line == "---":
                skip_title_block = False
            continue

        if not line:
            if in_table:
                add_table(doc, table_rows)
                table_rows = []
                in_table = False
            continue
        if line == "---":
            continue

        if line.startswith("|"):
            cols = [c.strip() for c in line.strip("|").split("|")]
            if all(set(c) <= {"-", ":"} for c in cols):
                continue
            table_rows.append(cols)
            in_table = True
            continue
        if in_table:
            add_table(doc, table_rows)
            table_rows = []
            in_table = False

        if line.startswith("## "):
            doc.add_heading(clean(line[3:]), level=1)
        elif line.startswith("### "):
            doc.add_heading(clean(line[4:]), level=2)
        elif line.startswith("- "):
            add_bullet(doc, line[2:])
        elif re.match(r"^\d+\.\s", line):
            add_number(doc, re.sub(r"^\d+\.\s", "", line))
        elif line.startswith("> "):
            add_quote(doc, line[2:])
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.1
            parse_inline(p, line)

    if in_table:
        add_table(doc, table_rows)

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build()
