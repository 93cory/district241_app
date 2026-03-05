"""
Générateur du document Word PNPI v2.0
Plateforme Nationale de Pilotage Industriel - Gabon
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Couleurs institutionnelles Gabon ────────────────────────────────────────
VERT_GABON    = RGBColor(0x00, 0x94, 0x40)   # vert drapeau
BLEU_GABON    = RGBColor(0x00, 0x3F, 0x8F)   # bleu institutionnel
JAUNE_GABON   = RGBColor(0xFF, 0xCD, 0x00)   # jaune drapeau
GRIS_CLAIR    = RGBColor(0xF4, 0xF4, 0xF4)
GRIS_TITRE    = RGBColor(0x1A, 0x1A, 0x2E)
BLANC         = RGBColor(0xFF, 0xFF, 0xFF)
ROUGE_ALERTE  = RGBColor(0xC0, 0x39, 0x2B)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), val.get('val', 'single'))
            b.set(qn('w:sz'), str(val.get('sz', 4)))
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(b)
    tcPr.append(tcBorders)

def add_horizontal_rule(doc, color='009440', width=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), str(width))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pb.append(bot)
    pPr.append(pb)
    return p

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def style_header_row(table, row_idx=0, bg='003F8F', text_color=BLANC):
    row = table.rows[row_idx]
    for cell in row.cells:
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = text_color
                run.font.size = Pt(9)

# ─── Document setup ──────────────────────────────────────────────────────────
doc = Document()
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ─── Default paragraph style ─────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name  = 'Calibri'
style.font.size  = Pt(10.5)
style.paragraph_format.space_after  = Pt(6)
style.paragraph_format.space_before = Pt(2)

# ─── Helper: colored run ─────────────────────────────────────────────────────
def colored_run(para, text, bold=False, color=None, size=None, italic=False):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color: run.font.color.rgb = color
    if size:  run.font.size = Pt(size)
    return run

def add_para(doc, text='', style='Normal', align=WD_ALIGN_PARAGRAPH.LEFT,
             before=2, after=6, bold=False, color=None, size=None):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        if color: run.font.color.rgb = color
        if size:  run.font.size = Pt(size)
    return p

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = VERT_GABON
    add_horizontal_rule(doc, color='009440', width=8)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLEU_GABON
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRIS_TITRE
    return p

def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.color.rgb = BLEU_GABON
        p.add_run(' ' + text)
    else:
        p.add_run(text)
    return p

def add_callout(doc, text, bg='F0F7F0', border_color='009440'):
    """Encadré coloré style callout."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.width = Cm(16)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run(text)
    run.bold = False
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    doc.add_paragraph()
    return table

def add_kpi_box(doc, items):
    """Boîte de KPIs en ligne (liste de (label, valeur, couleur_hex))."""
    n = len(items)
    table = doc.add_table(rows=2, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value, color) in enumerate(items):
        # valeur
        c_val = table.cell(0, i)
        set_cell_bg(c_val, color)
        p_val = c_val.paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_val.add_run(value)
        r.bold = True; r.font.size = Pt(18); r.font.color.rgb = BLANC
        c_val._tc.get_or_add_tcPr()
        # label
        c_lbl = table.cell(1, i)
        set_cell_bg(c_lbl, 'F4F4F4')
        p_lbl = c_lbl.paragraphs[0]
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p_lbl.add_run(label)
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = GRIS_TITRE
    doc.add_paragraph()
    return table

def simple_table(doc, headers, rows, col_widths=None, header_bg='003F8F'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # headers
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = BLANC
    # data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = 'FFFFFF' if r_idx % 2 == 0 else 'F4F7FB'
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            # bold first column
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            if c_idx == 0:
                run.bold = True
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE DE COUVERTURE
# ══════════════════════════════════════════════════════════════════════════════
# Bandeau vert supérieur
band = doc.add_table(rows=1, cols=1)
band.alignment = WD_TABLE_ALIGNMENT.CENTER
c = band.cell(0,0)
set_cell_bg(c, '009440')
c.width = Cm(16)
p = c.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(8)
r = p.add_run('RÉPUBLIQUE GABONAISE')
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = BLANC

doc.add_paragraph()

p = add_para(doc, 'MINISTÈRE DE L\'INDUSTRIE ET DE LA TRANSFORMATION LOCALE',
             align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=BLEU_GABON, size=12, before=6, after=4)

add_horizontal_rule(doc, color='FFCD00', width=12)

# Titre principal
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after  = Pt(6)
r = p.add_run('NOTE CONCEPTUELLE OFFICIELLE')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = GRIS_TITRE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(4)
r2 = p2.add_run('Plateforme Nationale de Pilotage Industriel')
r2.bold = True; r2.font.size = Pt(17); r2.font.color.rgb = VERT_GABON

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(20)
r3 = p3.add_run('PNPI — Volet Industrie')
r3.bold = False; r3.font.size = Pt(14); r3.font.color.rgb = BLEU_GABON

add_horizontal_rule(doc, color='009440', width=6)

# Boîte métadonnées
meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ('Version', '2.0 — Édition Stratégique'),
    ('Date', '26 février 2026'),
    ('Classification', 'Document institutionnel — Diffusion restreinte'),
    ('Statut', 'Soumis à validation du COPIL'),
]
for i, (k, v) in enumerate(meta_data):
    c1, c2 = meta.cell(i,0), meta.cell(i,1)
    set_cell_bg(c1, 'F4F4F4')
    set_cell_bg(c2, 'FFFFFF')
    c1.width = Cm(5); c2.width = Cm(11)
    r1 = c1.paragraphs[0].add_run(k)
    r1.bold = True; r1.font.size = Pt(9.5); r1.font.color.rgb = BLEU_GABON
    r2 = c2.paragraphs[0].add_run(v)
    r2.font.size = Pt(9.5)

doc.add_paragraph()

# Citation forte
add_callout(doc,
    '« Le PNPI n\'est pas un projet informatique. C\'est un instrument de politique '
    'industrielle et de souveraineté économique gabonaise. »',
    bg='E8F4EC', border_color='009440')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  RÉSUMÉ EXÉCUTIF
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, 'RÉSUMÉ EXÉCUTIF')

add_para(doc,
    'Le Gabon dispose d\'une base industrielle parmi les plus structurées d\'Afrique centrale, '
    'mais son administration industrielle reste l\'une des plus fragmentées de la sous-région. '
    'Le délai moyen de traitement d\'un dossier d\'Agrément Technique Industriel (ATI) dépasse '
    '45 à 60 jours, contre moins de 15 jours dans les pays comparables ayant digitalisé leurs '
    'procédures (Rwanda, Sénégal, Maroc).',
    before=4, after=6)

add_h2(doc, 'Ce que le PNPI apporte concrètement')

bullets_exec = [
    ('Réduction des délais ATI', '60 jours → moins de 15 jours (–75%)'),
    ('Traçabilité', '100% de chaque action sur chaque dossier, dès le premier jour'),
    ('Pilotage ministériel', 'Tableau de bord temps réel — le Ministre pilote l\'industrie nationale depuis son bureau'),
    ('Anti-corruption structurel', 'Élimination des contacts discrétionnaires dans le processus d\'agrément'),
    ('Leadership régional', 'Gabon positionné comme leader CEMAC de la gouvernance industrielle numérique'),
]
for label, desc in bullets_exec:
    add_bullet(doc, desc, bold_prefix=label + ' :')

doc.add_paragraph()

# KPI boxes
add_kpi_box(doc, [
    ('Réduction délai ATI',   '–75%',      '009440'),
    ('Dossiers complets (cible)', '>80%',   '003F8F'),
    ('Payback period',        '< 3 ans',   'FFCD00'.replace('CD00','A800')),
    ('Traçabilité',           '100%',      '1A1A2E'),
    ('ROI sur 5 ans',         '+145 M',    '009440'),
])

# Ligne budget
invest_table = doc.add_table(rows=1, cols=3)
invest_table.alignment = WD_TABLE_ALIGNMENT.CENTER
data_invest = [
    ('CAPEX\n(mise en place)', '330–360 M FCFA', '003F8F'),
    ('OPEX annuel\n(fonctionnement)', '110–130 M FCFA/an', '009440'),
    ('Payback period\n(retour sur investissement)', '< 3 ans', '1A2E6F'),
]
for i, (lbl, val, col) in enumerate(data_invest):
    cell = invest_table.cell(0, i)
    set_cell_bg(cell, col)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    r_val = p.add_run(val + '\n')
    r_val.bold = True; r_val.font.size = Pt(13); r_val.font.color.rgb = BLANC
    r_lbl = p.add_run(lbl)
    r_lbl.font.size = Pt(8); r_lbl.font.color.rgb = BLANC

doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. OBJET ET PÉRIMÈTRE
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '1. Objet et Périmètre')

add_h2(doc, '1.1 Objet')
add_para(doc,
    'La présente note définit le cadre institutionnel, technique, organisationnel, financier '
    'et stratégique de la Plateforme Nationale de Pilotage Industriel (PNPI), système '
    'd\'information souverain destiné à :',
    after=4)

obj_list = [
    ('Dématérialiser', 'les démarches prioritaires d\'agrément et de suivi industriel'),
    ('Centraliser', 'l\'ensemble des dossiers industriels dans un registre national unique'),
    ('Outiller', 'le Ministère de l\'Industrie d\'instruments de pilotage fiables et en temps réel'),
    ('Éliminer', 'les processus opaques et discrétionnaires qui nuisent à l\'attractivité du Gabon'),
    ('Préparer', 'l\'interopérabilité numérique avec ANPI, Douanes, DGI et Ministère du Commerce'),
]
for label, desc in obj_list:
    add_bullet(doc, desc, bold_prefix=label)

add_h2(doc, '1.2 Positionnement stratégique')
add_para(doc, 'Le PNPI s\'inscrit dans trois cadres stratégiques convergents :', after=4)
add_bullet(doc, 'Modernisation de l\'État, lutte contre la corruption, reddition de comptes',
           bold_prefix='Agenda de la Transition (CTRI) :')
add_bullet(doc, 'Transformation numérique des services publics gabonais',
           bold_prefix='Stratégie Gabon Digital :')
add_bullet(doc, 'Diversification économique, création d\'emplois, réduction des importations',
           bold_prefix='Engagements de développement industriel :')

doc.add_paragraph()
add_callout(doc,
    '« Le PNPI matérialise la promesse de rupture de la Transition dans le secteur économique. '
    'C\'est l\'instrument le plus visible et le plus immédiat de la réforme administrative '
    'industrielle. »',
    bg='EBF5FF', border_color='003F8F')

# ══════════════════════════════════════════════════════════════════════════════
#  2. CONTEXTE, DIAGNOSTIC ET JUSTIFICATION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '2. Contexte, Diagnostic et Justification')

add_h2(doc, '2.1 État des lieux quantifié')
add_para(doc, 'L\'administration industrielle gabonaise présente aujourd\'hui les caractéristiques suivantes :', after=4)

simple_table(doc,
    ['Indicateur', 'Situation actuelle', 'Cible PNPI'],
    [
        ['Délai moyen de traitement ATI', '45–60 jours', '< 15 jours'],
        ['Taux de dossiers complets au 1er dépôt', 'Estimé à 30–40%', '> 80%'],
        ['Traçabilité des décisions', 'Nulle — registres papier', '100%'],
        ['Visibilité sur le parc industriel', 'Partielle, non consolidée', 'Temps réel'],
        ['Coordination inter-administrations', 'Zéro — visites séquentielles', 'Plateforme partagée'],
        ['Données pour la décision ministérielle', 'Rapports trimestriels tardifs', 'Dashboard permanent'],
    ],
    col_widths=[6.5, 4.5, 4.5],
    header_bg='009440'
)

add_h2(doc, '2.2 Les cinq dysfonctionnements structurels')

dysfunctions = [
    ('① Fragmentation des circuits',
     'Un dossier ATI passe successivement dans plusieurs bureaux sans traçabilité ni délai contractuel. '
     'L\'opérateur ne sait pas où en est son dossier.'),
    ('② Opacité et discrétionnaire',
     'L\'absence de processus formalisé crée des espaces de discrétion qui favorisent les demandes '
     'informelles d\'accélération, nuisant à l\'égalité de traitement.'),
    ('③ Perte de mémoire institutionnelle',
     'Les mutations de personnel entraînent des pertes de dossiers. Une entreprise présente '
     'depuis 15 ans peut se voir demander des documents déjà fournis plusieurs fois.'),
    ('④ Reporting non consolidé',
     'Le Ministre reçoit des rapports trimestriels manuels, souvent obsolètes. '
     'Aucun outil ne permet de piloter l\'activité industrielle nationale en temps réel.'),
    ('⑤ Multiplication non coordonnée des contrôles',
     'Un industriel peut recevoir en quelques semaines la visite non coordonnée de l\'inspection '
     'du travail, de l\'environnement, de l\'industrie et des douanes, générant des coûts cachés significatifs.'),
]
for title, desc in dysfunctions:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.3)
    r1 = p.add_run(title + ' — ')
    r1.bold = True; r1.font.color.rgb = ROUGE_ALERTE
    p.add_run(desc)

add_h2(doc, '2.3 Le coût du statu quo')

simple_table(doc,
    ['Poste de coût', 'Estimation annuelle', 'Statut'],
    [
        ['Agents dédiés processus ATI (10 agents)', '60 M FCFA/an', 'Mesuré'],
        ['Archives physiques, courriers, déplacements', '5–8 M FCFA/an', 'Estimé'],
        ['Contentieux administratifs liés aux retards', 'Non mesuré', 'Réel'],
        ['Valeur économique bloquée (200 dossiers × 30j × 300K/j)', '1,8 Mds FCFA/an', 'Calculé'],
        ['Investissements perdus (image pays)', 'Non quantifié', 'Significatif'],
    ],
    col_widths=[8, 4, 3.5],
    header_bg='C0392B'
)

add_callout(doc,
    'CONCLUSION : Le statu quo coûte entre 1,5 et 2 milliards FCFA par an à l\'économie gabonaise. '
    'Le PNPI représente un investissement de 350 M FCFA — soit moins de 6 mois du coût de l\'inaction.',
    bg='FFF3CD', border_color='FFCD00')

add_h2(doc, '2.4 La fenêtre d\'opportunité de la Transition')
bullets_transit = [
    'La Transition (CTRI) a déclaré la rupture avec les pratiques opaques de l\'ancien régime',
    'La modernisation administrative est un outil de légitimation et de crédibilisation de la Transition',
    'Les partenaires financiers (BM, BAD, UE) attendent des projets structurants pour mobiliser leurs financements',
    'La Côte d\'Ivoire, le Cameroun et le Congo lancent leurs propres plateformes industrielles numériques — le Gabon ne peut rester en retrait',
]
for b in bullets_transit:
    add_bullet(doc, b)

# ══════════════════════════════════════════════════════════════════════════════
#  3. BENCHMARKS AFRICAINS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '3. Benchmarks Africains de Référence')

add_para(doc,
    'Les pays ci-dessous ont réalisé des transformations comparables au PNPI dans des contextes '
    'similaires. Leur expérience démontre la faisabilité et le retour sur investissement.',
    after=8)

benchmarks = [
    ('Rwanda — Irembo', '003F8F',
     'RÉFÉRENCE ABSOLUE',
     [
        '100+ services gouvernementaux intégrés',
        'Délai agrément industriel : 28 jours → 3 jours',
        'Satisfaction usagers : 95%',
        'Réduction coûts administratifs : 40%',
        'ROI atteint en 18 mois',
        'Budget : ~1,5 Mds FCFA',
        'Rang Doing Business : Top 40 africain',
     ],
     'Si le Rwanda, sans ressources naturelles, atteint le Top 40 grâce à la digitalisation, '
     'le Gabon peut atteindre le Top 20.'),
    ('Sénégal — APIX/CFE', '009440',
     'MODÈLE FRANCOPHONE',
     [
        'Contexte administratif francophone identique au Gabon',
        'Création d\'entreprise : 58 jours → 1 jour',
        'Budget comparable : ~800 M FCFA',
        'Déploiement en 14 mois',
        'Agréments industriels intégrés dans le Plan Sénégal Émergent',
     ],
     'Un système équivalent au PNPI peut être opérationnel en 14 mois dans un contexte '
     'administratif francophone identique.'),
    ('Maroc — CRI Digital', '1A2E6F',
     'RÉFÉRENCE INTEROPÉRABILITÉ',
     [
        'Interopérabilité avec 17 administrations en temps réel',
        'Zone industrielle : 18 mois → 4 mois de procédures',
        'API gouvernementale : Industrie + DGI + Douanes + Cadastre',
        'Certification industrielle 100% en ligne',
     ],
     'L\'interopérabilité n\'est pas une promesse lointaine. Le Maroc l\'a réalisée en 3 ans. '
     'Le PNPI pose les fondations pour que le Gabon suive ce chemin.'),
]

for title, color, badge, items, lesson in benchmarks:
    # Titre avec badge
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(f'  {title}  ')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = BLANC
    # workaround: use table for colored title
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    c1 = t.cell(0,0); c2 = t.cell(0,1)
    c1.width = Cm(9); c2.width = Cm(6.5)
    set_cell_bg(c1, color); set_cell_bg(c2, 'F4F4F4')
    pr = c1.paragraphs[0].add_run(title)
    pr.bold = True; pr.font.size = Pt(11); pr.font.color.rgb = BLANC
    c1.paragraphs[0].paragraph_format.space_before = Pt(4)
    c1.paragraphs[0].paragraph_format.space_after  = Pt(4)
    pr2 = c2.paragraphs[0].add_run(badge)
    pr2.bold = True; pr2.font.size = Pt(8.5)
    pr2.font.color.rgb = RGBColor(*bytes.fromhex(color))
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c2.paragraphs[0].paragraph_format.space_before = Pt(4)
    for item in items:
        add_bullet(doc, item, level=0)
    add_callout(doc, f'Leçon pour le Gabon : {lesson}', bg='E8F4EC')

# Tableau comparatif synthétique
add_h2(doc, 'Tableau comparatif synthétique')
simple_table(doc,
    ['Pays', 'Plateforme', 'Budget', 'Délai', 'Impact mesuré'],
    [
        ['Rwanda',       'Irembo',    '1,5 Mds FCFA',  '24 mois', 'Délai industriel : 28j→3j'],
        ['Sénégal',      'APIX/CFE',  '~800 M FCFA',   '14 mois', 'Création ent. : 58j→1j'],
        ['Maroc',        'CRI Digital','N.D.',          '36 mois', 'Zone ind. : 18m→4m'],
        ['Côte d\'Ivoire','CEPICI',   'N.D.',           '18 mois', '–30% délai traitement'],
        ['GABON (PNPI)', 'PNPI',      '340–360 M FCFA','12 mois', 'Cible : ATI 60j→15j'],
    ],
    col_widths=[3, 3, 3.5, 2.5, 4],
    header_bg='009440'
)

# ══════════════════════════════════════════════════════════════════════════════
#  4. OBJECTIFS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '4. Objectifs')

add_h2(doc, '4.1 Objectif général')
add_para(doc,
    'Mettre en place un système d\'information industriel souverain, sécurisé et interopérable, '
    'permettant au Ministère de l\'Industrie de piloter l\'ensemble des processus d\'agrément, '
    'de suivi et de conformité industrielle en temps réel, tout en offrant aux opérateurs '
    'économiques une expérience administrative de classe mondiale.',
    after=8)

add_h2(doc, '4.2 Objectifs spécifiques')
simple_table(doc,
    ['#', 'Objectif', 'Mesure de succès'],
    [
        ['1', 'Dématérialiser 100% des démarches ATI et dossiers prioritaires', '0 dossier traité hors plateforme à M12'],
        ['2', 'Réduire le délai ATI de 60 à < 15 jours', 'Délai mesuré mensuellement sur le dashboard'],
        ['3', 'Atteindre > 80% de dossiers complets au 1er dépôt', 'Taux calculé automatiquement'],
        ['4', 'Fournir un tableau de bord temps réel au Ministre', 'Dashboard opérationnel dès M7'],
        ['5', 'Tracer 100% des actions sur les dossiers critiques', 'Audit log exhaustif dès J1'],
        ['6', 'Constituer le Dossier Industriel Unique (DIU)', 'Registre complet à M12'],
        ['7', 'Préparer l\'interopérabilité ANPI, DGI, Douanes', 'API Gateway avec stubs à M12'],
        ['8', 'Éliminer les contacts discrétionnaires', 'Zéro étape non tracée dans le workflow'],
    ],
    col_widths=[1, 8.5, 6],
    header_bg='003F8F'
)

# ══════════════════════════════════════════════════════════════════════════════
#  5. MODULES DÉTAILLÉS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '5. Périmètre Phase 1 — Modules Détaillés')

modules = [
    ('Module 1', 'Guichet Numérique Unique Industriel', '003F8F',
     'Portail de dépôt et de suivi accessible aux opérateurs industriels via navigateur web et mobile.',
     [
        'Création de compte opérateur avec vérification d\'identité (NIFGABON, registre ANPI)',
        'Formulaires intelligents adaptés au type d\'activité avec guidage pas à pas',
        'Vérification automatique de complétude en temps réel',
        'Numéro de dossier unique et accusé de réception électronique',
        'Suivi d\'avancement par étape (tracking temps réel)',
        'Notifications automatiques par email et SMS',
        'Mode offline pour zones à connectivité limitée',
     ],
     'Un opérateur de Franceville ou Port-Gentil dépose son dossier sans se déplacer à Libreville. '
     'Le taux de dossiers incomplets passe de ~35% à > 80%.'),

    ('Module 2', 'Workflow d\'Instruction et de Validation', '009440',
     'Cœur de la plateforme côté Ministère — circuit complet d\'instruction d\'un dossier.',
     [
        'Attribution automatique des dossiers selon la charge de travail',
        'Séquençage : pré-instruction → technique → juridique → décision',
        'Délais réglementaires configurables par type de dossier (SLA)',
        'Alertes automatiques pour dossiers proches de la limite délai',
        'Escalade automatique au supérieur si délai dépassé',
        'Signature électronique multi-niveaux (agent → chef service → Directeur)',
        'Génération automatique des décisions en PDF signés',
     ],
     'Fin de l\'arbitraire. Chaque dossier suit le même chemin dans les mêmes délais.'),

    ('Module 3', 'Gestion des ATI (Agréments Techniques Industriels)', 'C0392B',
     'Module spécialisé pour le processus ATI avec ses spécificités réglementaires.',
     [
        'Formulaires ATI par secteur : bois/forêt, mines, agroalimentaire, BTP',
        'Vérification automatique de cohérence des données déclarées',
        'Circuit multi-niveaux (Direction Technique → Juridique → Secrétariat Général)',
        'Génération de l\'ATI numérique avec QR code d\'authenticité',
        'Registre public de vérification des ATI (banques, partenaires)',
        'Alertes de renouvellement automatiques (90/60/30 jours avant expiration)',
        'Historique complet (modifications, renouvellements, suspensions)',
     ],
     'Élimination de la falsification des ATI. Un QR code suffit pour vérifier l\'authenticité.'),

    ('Module 4', 'Dossier Industriel Unique (DIU) et Registre Central', '1A2E6F',
     'Mémoire institutionnelle permanente de chaque opérateur industriel au Gabon.',
     [
        'Fiche permanente par entreprise (identité, historique, agréments actifs)',
        'Consolidation automatique de tous les documents — jamais perdus, jamais à refournir',
        'Moteur de recherche avancé (secteur, région, taille, statut)',
        'Base de référence pour les statistiques industrielles nationales',
        'Export des données pour les besoins décisionnels',
     ],
     'Fin de la perte de mémoire institutionnelle. Un nouvel agent a accès à tout l\'historique instantanément.'),

    ('Module 5', 'Dashboard Ministériel et Business Intelligence', '009440',
     'Tableau de bord décisionnel pour le Ministre, le SG, les Directeurs et leurs équipes.',
     [
        'Vue synthétique temps réel : dossiers, délais, taux de traitement',
        'Cartographie industrielle du Gabon (GIS) par secteur et province',
        'Indicateurs sectoriels : agréments par filière (bois, mines, agroalimentaire…)',
        'Suivi KPIs nationaux : emplois industriels, valeur ajoutée, taux de transformation',
        'Alertes décisionnelles : dossiers critiques, anomalies',
        'Rapports automatisés exportables (PDF/Excel)',
        'Prévisions et tendances basées sur les données historiques',
     ],
     'Le Ministre pilote pour la première fois le secteur industriel national avec des données fiables et actualisées.'),

    ('Module 6', 'Conformité Post-Agrément et Inspections Terrain', '7D3C98',
     'Suivi de la vie de l\'entreprise après l\'obtention de l\'agrément.',
     [
        'Calendrier d\'inspections programmées automatiquement',
        'Coordination inter-administrations (inspections conjointes)',
        'Application mobile inspecteurs : rapport géolocalisé, photos, signature',
        'Suivi des mesures correctives avec délais et alertes',
        'Tableau de bord conformité sectorielle par province',
     ],
     'L\'agrément n\'est plus un document qu\'on délivre et oublie. La conformité est suivie dans le temps.'),
]

for mod_num, mod_title, color, desc, items, value in modules:
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    c1 = t.cell(0,0); c2 = t.cell(0,1)
    c1.width = Cm(2.5); c2.width = Cm(13)
    set_cell_bg(c1, color)
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c1.paragraphs[0].paragraph_format.space_before = Pt(6)
    r = c1.paragraphs[0].add_run(mod_num)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = BLANC
    set_cell_bg(c2, 'F4F7FB')
    c2.paragraphs[0].paragraph_format.space_before = Pt(4)
    r2 = c2.paragraphs[0].add_run(mod_title)
    r2.bold = True; r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(*bytes.fromhex(color))

    add_para(doc, desc, before=4, after=3)
    for item in items:
        add_bullet(doc, item)
    add_callout(doc, f'Valeur ajoutée : {value}', bg='EBF5FF')

# ══════════════════════════════════════════════════════════════════════════════
#  6. FONCTIONNALITÉS CLÉS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '6. Fonctionnalités Clés et Valeur Ajoutée')

features = [
    ('6.1 Intelligence Artificielle — Pré-instruction automatique', [
        'Lecture et extraction automatique des informations des documents (PDF, scans) via OCR',
        'Vérification de cohérence entre données déclarées et bases de référence',
        'Calcul d\'un score de complétude et d\'un score de risque en < 2 minutes',
        'Génération d\'un rapport de pré-instruction pour l\'agent humain',
        'Technologie : LangChain + Mistral (open source français) + OCR Tesseract',
        'Budget additionnel : 15–20 M FCFA | Gain : 3–5 jours de vérification → 2 minutes',
    ]),
    ('6.2 Signature Électronique à Valeur Juridique (PKI)', [
        'ATI numériques ayant la même valeur juridique que les originaux papier',
        'Horodatage certifié et infalsifiable de chaque décision',
        'Vérification d\'authenticité instantanée par QR code pour les tiers (banques, partenaires)',
        'Élimination du marché des faux agréments et certifications',
        'Prérequis : décret donnant valeur juridique à la signature électronique industrielle',
    ]),
    ('6.3 Cartographie Industrielle Géospatiale (GIS)', [
        'Localisation de toutes les entreprises agréées par secteur, province et statut',
        'Visualisation des clusters industriels et déserts industriels',
        'Flux d\'agréments et de renouvellements dans le temps',
        'Corrélation agréments accordés / emplois créés (données CNSS)',
        'Outil de politique industrielle — la carte que le Gabon n\'a jamais eu',
    ]),
    ('6.4 API Gateway — Socle d\'Interopérabilité Future', [
        'Architecture permettant la connexion progressive avec les administrations partenaires',
        'ANPI-Gabon : vérification automatique du registre des entreprises',
        'DGI : vérification du statut fiscal des demandeurs',
        'Douanes (SYDONIA) : vérification des importations d\'équipements déclarés',
        'CNSS : vérification des déclarations sociales',
        'Dès la Phase 1 : stubs prêts, activation progressive par convention',
    ]),
]

for title, items in features:
    add_h2(doc, title)
    for item in items:
        add_bullet(doc, item)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  7. ARCHITECTURE TECHNIQUE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '7. Architecture Technique et Sécurité')

add_h2(doc, '7.1 Choix technologiques')
simple_table(doc,
    ['Composant', 'Technologie', 'Justification'],
    [
        ['Backend API', 'FastAPI (Python 3.12)', 'Performance, maturité, écosystème IA/ML, open source'],
        ['Base de données', 'PostgreSQL 16', 'Robustesse, ACID, support JSON, open source, souverain'],
        ['Frontend dashboard', 'Next.js 14 / React', 'Standard industrie, SSR, performance, maintenable'],
        ['Application mobile', 'Flutter', 'Multi-plateforme (iOS + Android) depuis un seul code'],
        ['Authentification', 'JWT + OAuth2 + MFA', 'Sécurité renforcée, standard international'],
        ['IA documentaire', 'LangChain + Mistral', 'Open source français — souveraineté numérique'],
        ['Conteneurisation', 'Docker + Compose', 'Déploiement reproductible et portable'],
        ['CI/CD', 'GitHub Actions', 'Automatisation tests et déploiements'],
        ['Monitoring', 'Prometheus + Grafana', 'Supervision temps réel, alertes'],
        ['Migrations BDD', 'Alembic', 'Versionnage rigoureux du schéma de données'],
    ],
    col_widths=[4, 4, 7.5],
    header_bg='003F8F'
)

add_h2(doc, '7.2 Hébergement souverain')
add_callout(doc,
    'PRINCIPE INVIOLABLE : Les données industrielles stratégiques du Gabon ne transitent '
    'sur aucun serveur étranger. Le PNPI est hébergé sur infrastructure souveraine gabonaise.',
    bg='FFF3CD', border_color='FFCD00')
add_bullet(doc, 'Datacenter du Gouvernement Gabonais (si disponible et certifié)', bold_prefix='Option 1 (recommandée) :')
add_bullet(doc, 'Cloud souverain africain (AfricaCloud, Dimension Data Gabon)', bold_prefix='Option 2 :')
add_bullet(doc, 'Infrastructure dédiée hébergée localement avec redondance', bold_prefix='Option 3 :')

add_h2(doc, '7.3 Sécurité — Architecture multi-couches')
simple_table(doc,
    ['Couche', 'Mesures de sécurité'],
    [
        ['Accès', 'HTTPS/TLS 1.3 — MFA obligatoire — Rate limiting — RBAC 5 niveaux'],
        ['Application', 'Validation entrées (injection SQL, XSS) — JWT rotation — Sessions limitées — AES-256'],
        ['Données', 'Audit log infalsifiable — Sauvegardes chiffrées quotidiennes — RPO < 24h / RTO < 4h'],
        ['Audits', 'Pentest externe annuel — Revue code sécurité avant production — Formation agents'],
    ],
    col_widths=[3.5, 12],
    header_bg='C0392B'
)

# ══════════════════════════════════════════════════════════════════════════════
#  8. ROI ET JUSTIFICATION ÉCONOMIQUE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '8. Retour sur Investissement et Justification Économique')

add_h2(doc, '8.1 Tableau de ROI sur 5 ans')
simple_table(doc,
    ['Poste', 'Valeur annuelle'],
    [
        ['CAPEX (mise en place)', '340 M FCFA (one-shot)'],
        ['OPEX annuel (fonctionnement)', '120 M FCFA/an'],
        ['Économies directes (masse salariale, archives, contentieux)', '31–37 M FCFA/an'],
        ['Valeur économique (accélération agréments, 10% attributabilité)', '180 M FCFA/an'],
        ['TOTAL VALEUR ANNUELLE (conservateur)', '211–217 M FCFA/an'],
        ['Payback period', '340 M / 217 M = 1,6 an après déploiement'],
        ['ROI net sur 5 ans (5 × 217M − 940M)', '+145 M FCFA minimum'],
    ],
    col_widths=[10, 5.5],
    header_bg='009440'
)

add_h2(doc, '8.2 Impact sur l\'attractivité internationale')
add_para(doc,
    'Le PNPI contribue directement à l\'amélioration de deux indicateurs internationaux mesurés :',
    after=4)
add_bullet(doc,
    'Indicateur "Obtention des licences d\'affaires" directement impacté. '
    'Chaque gain de 10 rangs génère 2–3% d\'IDE supplémentaires pour les pays d\'ASS. '
    'Pour le Gabon : 50–100 millions USD d\'investissements additionnels annuels potentiels.',
    bold_prefix='B-READY Banque Mondiale :')
add_bullet(doc,
    'Pilier "Efficacité des services publics" amélioré — signal positif pour les investisseurs institutionnels.',
    bold_prefix='Mo Ibrahim Index :')

add_callout(doc,
    'FORMULATION DÉCISIONNELLE : Le coût total du PNPI sur 5 ans est de 940 M FCFA. '
    'L\'impact potentiel sur l\'attraction des IDE représente 50 à 100 fois cet investissement '
    'annuellement. Ce n\'est pas une dépense publique. C\'est l\'investissement à plus fort '
    'effet de levier disponible aujourd\'hui pour l\'économie gabonaise.',
    bg='E8F4EC', border_color='009440')

add_h2(doc, '8.3 Modalités de financement envisagées')
simple_table(doc,
    ['Source de financement', 'Part estimée', 'Montant', 'Condition'],
    [
        ['Budget national', '40–50%', '~150 M FCFA', 'Décision COPIL'],
        ['Banque Mondiale / IFC', '20–30%', '~80 M FCFA', 'Dossier Investment Climate'],
        ['Union Européenne (PASE/PARCIP)', '15–20%', '~60 M FCFA', 'Éligibilité réforme administrative'],
        ['AFD (modernisation admin. francophone)', '10–15%', '~45 M FCFA', 'Convention France-Gabon'],
        ['BAD (Technologies pour l\'Afrique)', '10–15%', '~40 M FCFA', 'Dossier TTA soumis'],
    ],
    col_widths=[5.5, 2.5, 3, 4.5],
    header_bg='003F8F'
)
add_callout(doc,
    'La part du budget national peut être réduite à 150–170 M FCFA si les partenaires financiers '
    'sont mobilisés en parallèle. Le PNPI est un catalyseur de financement international, '
    'pas une charge budgétaire.',
    bg='EBF5FF', border_color='003F8F')

# ══════════════════════════════════════════════════════════════════════════════
#  9. BUDGET DÉTAILLÉ
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '9. Budget Détaillé')

add_h2(doc, '9.1 CAPEX — Décomposition poste par poste')
simple_table(doc,
    ['#', 'Poste', 'Description', 'Montant FCFA'],
    [
        ['1',  'Cadrage et conception',         'UX/UI, architecture, spécifications, maquettes',          '28–35 M'],
        ['2',  'Développement Backend',          'API FastAPI, BDD PostgreSQL, workflows, sécurité',        '65–80 M'],
        ['3',  'Frontend Dashboard',             'Next.js, tableaux de bord, cartographie GIS',             '45–55 M'],
        ['4',  'Application Mobile',             'Flutter (iOS + Android), mode offline, géolocalisation',  '28–35 M'],
        ['5',  'Module IA Documentaire',         'OCR + NLP pré-instruction automatique',                   '15–20 M'],
        ['6',  'Signature Électronique (PKI)',   'Infrastructure clés, intégration, certification',         '20–25 M'],
        ['7',  'Infrastructure et Hébergement',  'Serveurs, CDN, stockage, redondance, setup',              '42–55 M'],
        ['8',  'Intégrations et API Gateway',    'Stubs ANPI, DGI, Douanes + architecture interopérabilité','22–28 M'],
        ['9',  'Sécurité et Audit initial',      'Pentest, hardening, configuration MFA',                   '15–20 M'],
        ['10', 'Formation',                       '50+ agents, présentiel + e-learning',                     '18–22 M'],
        ['11', 'Conduite du Changement',         'Communication, accompagnement, sensibilisation PME',      '12–15 M'],
        ['12', 'Tests, Recette, Déploiement',    'UAT, tests de charge, go-live, hypercare',                '15–20 M'],
        ['13', 'Documentation et Transfert',     'Manuels, formation formateurs, code source',              '8–10 M'],
        ['14', 'Contingence (10%)',               'Aléas et imprévus',                                       '33–40 M'],
        ['',   'TOTAL CAPEX',                    'Scénario de référence : 340–360 M FCFA',                  '366–460 M'],
    ],
    col_widths=[0.8, 4.5, 6, 3],
    header_bg='003F8F'
)

add_h2(doc, '9.2 OPEX annuel')
simple_table(doc,
    ['Poste', 'Description', 'Montant/an'],
    [
        ['Maintenance applicative', '2–3 développeurs demi-temps', '40–50 M'],
        ['Hébergement et infrastructure', 'Datacenter, licences, CDN', '20–30 M'],
        ['Support utilisateurs', 'Hotline N1/N2, assistance agents', '15–20 M'],
        ['Évolutions fonctionnelles', 'Nouvelles fonctionnalités, adaptations', '20–25 M'],
        ['Sécurité continue', 'Pentest annuel, veille, mises à jour', '10–15 M'],
        ['Formation continue', 'Nouveaux agents, nouvelles fonctionnalités', '5–8 M'],
        ['TOTAL OPEX', 'Scénario de référence : 120–130 M FCFA/an', '110–148 M'],
    ],
    col_widths=[5, 7, 3.5],
    header_bg='009440'
)

add_h2(doc, '9.3 Validation du budget par les comparatifs de marché')
add_bullet(doc, '1,5 Mds FCFA pour 100 services → PNPI est ~4x moins cher pour un périmètre spécialisé. ✅', bold_prefix='Irembo Rwanda :')
add_bullet(doc, '~800 M FCFA périmètre plus large → PNPI ~50% moins cher. ✅', bold_prefix='CFE Sénégal :')
add_bullet(doc, '600 M–1 Mds FCFA pour équivalent → PNPI 40–50% d\'économie. ✅', bold_prefix='Cabinet international (Deloitte/Capgemini) :')

# ══════════════════════════════════════════════════════════════════════════════
#  10. PLAN DE MISE EN ŒUVRE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '10. Plan de Mise en Œuvre (12 mois)')

simple_table(doc,
    ['Phase', 'Mois', 'Activités principales', 'Livrable clé'],
    [
        ['M1–M2\nCADRAGE',      'Mois 1–2',   'Spécifications, UX/UI, architecture, constitution équipe',      'Cahier des charges validé + maquettes'],
        ['M3–M5\nBUILD MVP',    'Mois 3–5',   'Développement Modules 1–3 (Guichet, Workflow ATI, DIU)',         'MVP fonctionnel + tests'],
        ['M6–M7\nPILOTE',       'Mois 6–7',   'Pilote opérationnel ATI + dashboard avec données réelles',      '50 dossiers ATI traités sur la plateforme'],
        ['M8–M10\nSTABILISATION','Mois 8–10', 'Modules 4–6 (BI, Conformité, Admin), montée en charge',         'Plateforme complète opérationnelle'],
        ['M11–M12\nEXTENSION',  'Mois 11–12', 'Formation complète, transfert de compétences, bilan, V2 planifiée','100% dossiers ATI sur la plateforme'],
    ],
    col_widths=[3, 2, 7, 4],
    header_bg='009440'
)

add_h2(doc, 'Critères de passage de phase')
phases = [
    ('M2→M3', [
        'Cahier des charges signé',
        'Équipe projet constituée et opérationnelle',
        'Infrastructure de développement opérationnelle',
        'Budget validé et engagé',
    ]),
    ('M5→M6', [
        'MVP démontré et validé par les référents métier',
        'Formation des agents pilotes réalisée',
        '0 anomalie bloquante en recette fonctionnelle',
    ]),
    ('M7→M8', [
        '50 dossiers ATI traités sans incident majeur',
        'KPIs pilote atteints : délai < 20 jours, dossiers complets > 70%',
        'Satisfaction agents instructeurs > 70%',
    ]),
    ('M12 — Validation finale', [
        '100% des nouveaux dossiers ATI via PNPI',
        'Disponibilité applicative > 99,5%',
        'Formation 100% agents complétée',
        'PCA/PRA testés et validés',
    ]),
]
for phase_name, criteria in phases:
    add_h3(doc, f'Passage {phase_name}')
    for c in criteria:
        add_bullet(doc, c)

# ══════════════════════════════════════════════════════════════════════════════
#  11. KPIs
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '11. Indicateurs de Performance (KPIs)')

add_h2(doc, 'KPIs Projet — Déploiement')
simple_table(doc,
    ['KPI', 'Cible M7 (pilote)', 'Cible M12 (généralisation)'],
    [
        ['% agents formés',              '80% (agents pilotes)', '100%'],
        ['% dossiers ATI sur plateforme', '50%',                 '100%'],
        ['Disponibilité applicative',    '> 99%',               '> 99,5%'],
        ['Incidents bloquants ouverts',  '< 5',                 '0'],
    ],
    col_widths=[7, 4, 4.5],
    header_bg='009440'
)

add_h2(doc, 'KPIs Métier — Performance')
simple_table(doc,
    ['KPI', 'Baseline actuelle', 'Cible M7', 'Cible M12'],
    [
        ['Délai moyen traitement ATI',         '45–60 jours',    '< 25 jours', '< 15 jours'],
        ['Dossiers complets au 1er dépôt',     '~35%',           '> 65%',      '> 80%'],
        ['Traitement dans les délais SLA',     'Non mesuré',     '> 70%',      '> 90%'],
        ['Traçabilité des décisions',          '0%',             '100%',       '100%'],
        ['Satisfaction opérateurs',            'Non mesurée',    '> 65%',      '> 80%'],
        ['Satisfaction agents',                'Non mesurée',    '> 60%',      '> 75%'],
    ],
    col_widths=[6, 3, 2.5, 3],
    header_bg='003F8F'
)

add_h2(doc, 'KPIs Stratégiques — Impact')
simple_table(doc,
    ['KPI', 'Source', 'Fréquence'],
    [
        ['Évolution rang B-READY (indicateur licences)', 'Banque Mondiale', 'Annuelle'],
        ['Nombre d\'ATI traités et validés',             'PNPI',            'Mensuelle'],
        ['Économies générées vs statu quo',              'Audit interne',   'Semestrielle'],
        ['Connexions interopérabilité activées',         'Équipe technique', 'Trimestrielle'],
        ['Taux d\'adoption utilisateurs',                'PNPI',            'Mensuelle'],
    ],
    col_widths=[8, 4, 3.5],
    header_bg='1A2E6F'
)

# ══════════════════════════════════════════════════════════════════════════════
#  12. GOUVERNANCE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '12. Gouvernance')

add_h2(doc, '12.1 Fréquence des instances')
simple_table(doc,
    ['Instance', 'Fréquence', 'Participants', 'Objet'],
    [
        ['COPIL Stratégique',    'Mensuel',        'Ministre/SG, Direction Programme, Lead Technique', 'Arbitrage stratégique, KPIs'],
        ['Comité Opérationnel',  'Hebdomadaire',   'Direction Programme, PO, Lead Technique',          'Avancement, blocages'],
        ['Revue Technique',      'Bi-hebdomadaire','Lead Technique, DevOps, QA',                       'Architecture, code, tests'],
        ['Revue KPI',            'Mensuelle',      'Direction Programme, PO, Data Analyst',            'Indicateurs, ajustements'],
    ],
    col_widths=[3.5, 2.5, 5.5, 4],
    header_bg='003F8F'
)

add_h2(doc, '12.2 Modèle RACI')
simple_table(doc,
    ['Activité', 'Dir. Programme', 'Product Owner', 'Lead Technique', 'COPIL'],
    [
        ['Cadrage et spécifications', 'A', 'R', 'C', 'I'],
        ['Développement',             'C', 'I', 'A+R', 'I'],
        ['Recette fonctionnelle',     'C', 'A+R', 'C', 'I'],
        ['Mise en production',        'I', 'I', 'A+R', 'I'],
        ['Formation agents',          'A', 'R', 'C', 'I'],
        ['Décisions budgétaires',     'C', 'C', 'I', 'A'],
    ],
    col_widths=[5.5, 2.5, 2.5, 2.5, 2.5],
    header_bg='009440'
)
add_para(doc, 'R = Réalisateur | A = Approbateur | C = Consulté | I = Informé', size=8.5, before=2, after=6)

# ══════════════════════════════════════════════════════════════════════════════
#  13. CADRE LÉGAL
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '13. Cadre Légal et Conformité')

add_h2(doc, '13.1 Textes applicables')
legal_texts = [
    ('Loi n°009/2012', 'Code des investissements du Gabon — cadre des agréments industriels'),
    ('Loi n°025/2018', 'Protection des données personnelles au Gabon'),
    ('Stratégie Gabon Digital', 'Référentiel national de transformation numérique'),
    ('Accords CEMAC', 'Harmonisation des procédures industrielles en Afrique centrale'),
    ('RGPD', 'Applicable aux entreprises européennes opérant au Gabon'),
]
for ref, desc in legal_texts:
    add_bullet(doc, desc, bold_prefix=ref + ' :')

add_h2(doc, '13.2 Travaux réglementaires à mener en parallèle')
simple_table(doc,
    ['Texte nécessaire', 'Objet', 'Priorité'],
    [
        ['Décret ou arrêté', 'Valeur juridique des agréments électroniques', 'CRITIQUE — avant M6'],
        ['Arrêté ministériel', 'Définition et valeur du Dossier Industriel Unique (DIU)', 'Haute — avant M8'],
        ['Convention inter-admin.', 'Cadre interopérabilité ANPI, DGI, Douanes', 'Moyenne — avant M10'],
        ['Circulaire', 'Obligation d\'utilisation du PNPI pour tous les dépôts', 'CRITIQUE — avant M12'],
    ],
    col_widths=[4, 7, 3.5],
    header_bg='C0392B'
)

# ══════════════════════════════════════════════════════════════════════════════
#  14. MATRICE DES RISQUES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '14. Matrice des Risques')

simple_table(doc,
    ['#', 'Risque', 'Prob.', 'Impact', 'Niveau', 'Mesure d\'atténuation'],
    [
        ['R1', 'Résistance au changement des agents',    'Haute',   'Critique', 'CRITIQUE', 'Plan change management dès M1, sponsors métiers, référents internes'],
        ['R2', 'Faible adoption des opérateurs',          'Moyenne', 'Majeur',   'ÉLEVÉ',    'Module assistance PME, accompagnement terrain, communication ciblée'],
        ['R3', 'Budget insuffisant ou non engagé',        'Moyenne', 'Critique', 'ÉLEVÉ',    'Phasage modulaire, co-financements extérieurs en parallèle'],
        ['R4', 'Connectivité insuffisante en provinces',  'Haute',   'Majeur',   'ÉLEVÉ',    'Mode offline-first, synchronisation différée, VSAT zones critiques'],
        ['R5', 'Changement politique ou administratif',   'Faible',  'Critique', 'ÉLEVÉ',    'Ancrage multi-acteurs, documentation exhaustive, transfert compétences'],
        ['R6', 'Risque cybersécurité / intrusion',        'Faible',  'Critique', 'MOYEN',    'MFA, audit, journalisation exhaustive, PRA/PCA, pentest régulier'],
        ['R7', 'Qualité insuffisante des données initiales','Haute', 'Modéré',   'MOYEN',    'Règles de validation strictes, module de migration des données'],
        ['R8', 'Dépendance fournisseur technologique',    'Moyenne', 'Modéré',   'MOYEN',    'Clauses de réversibilité, code open source, transfert compétences'],
        ['R9', 'Délai de mise en œuvre dépassé',          'Moyenne', 'Modéré',   'MOYEN',    'Méthode Agile, jalons intermédiaires, contingence 10%'],
        ['R10','Interopérabilité bloquée',                'Haute',   'Faible',   'FAIBLE',   'Architecture stubs — activation progressive — Phase 1 indépendante'],
    ],
    col_widths=[0.8, 4.5, 1.5, 1.8, 1.8, 5.5],
    header_bg='C0392B'
)

# ══════════════════════════════════════════════════════════════════════════════
#  15. DÉPLOIEMENT TERRITORIAL
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '15. Stratégie de Déploiement Territorial')

add_h2(doc, '15.1 Approche par vagues (9 provinces)')
simple_table(doc,
    ['Vague', 'Période', 'Couverture', 'Justification'],
    [
        ['Vague 1', 'M6–M8',   'Libreville (Estuaire)',               '80% des dossiers ATI — connectivité fibre — sièges sociaux'],
        ['Vague 2', 'M9–M10',  'Port-Gentil + Franceville',            'Industrie pétrolière (Ogooué-Maritime) + minière (Comilog)'],
        ['Vague 3', 'M11–M12', 'Oyem, Lambaréné, Mouila + provinces', 'Agroalimentaire, forêt, pêche — déploiement national complet'],
    ],
    col_widths=[2, 2, 5, 6.5],
    header_bg='009440'
)

add_h2(doc, '15.2 Solutions pour les zones à connectivité limitée')
add_bullet(doc, 'L\'application mobile des inspecteurs fonctionne sans connexion. Données synchronisées dès reconnexion.', bold_prefix='Mode offline-first :')
add_bullet(doc, 'Points d\'accès dédiés dans chaque chef-lieu provincial (mairies, agences ANPI locales).', bold_prefix='Points d\'accès dédiés :')
add_bullet(doc, 'Formation d\'agents-relais provinciaux capables d\'assister les opérateurs au dépôt numérique.', bold_prefix='Agents-relais :')

# ══════════════════════════════════════════════════════════════════════════════
#  16. DÉCISIONS ATTENDUES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '16. Décisions Attendues')

add_para(doc,
    'À l\'issue de la présentation de cette note, les décisions suivantes sont sollicitées :',
    after=6)

simple_table(doc,
    ['#', 'Décision', 'Responsable', 'Délai souhaité'],
    [
        ['1', 'Validation du lancement de la Phase 1',                  'Ministre / COPIL',         'J+15'],
        ['2', 'Validation de la gouvernance proposée',                  'Secrétaire Général',        'J+15'],
        ['3', 'Validation de l\'enveloppe budgétaire',                  'Ministère des Finances',   'J+30'],
        ['4', 'Désignation du Directeur de Programme',                  'Ministre',                  'J+10'],
        ['5', 'Désignation des référents métiers',                      'Directeurs Généraux',       'J+15'],
        ['6', 'Lancement des démarches de co-financement (BM/UE/AFD)',  'Direction Programme',       'J+30'],
        ['7', 'Lancement des travaux réglementaires (valeur juridique ATI numérique)', 'Dir. Juridique', 'J+30'],
    ],
    col_widths=[0.7, 7, 4, 2.5],
    header_bg='003F8F'
)

# ══════════════════════════════════════════════════════════════════════════════
#  CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, 'Conclusion')

add_para(doc,
    'Le PNPI n\'est pas un projet de modernisation administrative parmi d\'autres. '
    'C\'est l\'instrument qui transforme le Ministère de l\'Industrie en institution '
    'de pilotage industriel de classe mondiale.',
    bold=False, before=6, after=8)

conclusions = [
    'Réduit de 75% les délais d\'agrément industriel en 12 mois',
    'Élimine structurellement la corruption dans les processus d\'agrément',
    'Dote le Gabon d\'une cartographie industrielle nationale — pour la première fois de son histoire',
    'Affirme la souveraineté numérique gabonaise sur ses données industrielles stratégiques',
    'Rejoint les nations africaines leaders de la gouvernance économique digitale',
]
for c in conclusions:
    add_bullet(doc, c)

doc.add_paragraph()

add_callout(doc,
    'Les pays qui nous précèdent — Rwanda, Sénégal, Maroc, Côte d\'Ivoire — ont prouvé que c\'est '
    'possible. Le contexte de la Transition en fait le moment idéal. Les partenaires financiers '
    'sont prêts.\n\n'
    'Il ne manque que la décision.',
    bg='E8F4EC', border_color='009440')

doc.add_paragraph()

# Bandeau final
band2 = doc.add_table(rows=1, cols=3)
band2.alignment = WD_TABLE_ALIGNMENT.CENTER
colors_band = ['009440', 'FFCD00', '003F8F']
labels_band = ['Souveraineté', 'Transparence', 'Performance']
for i, (col, lbl) in enumerate(zip(colors_band, labels_band)):
    c = band2.cell(0, i)
    set_cell_bg(c, col)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(lbl)
    r.bold = True; r.font.size = Pt(12)
    r.font.color.rgb = BLANC if col != 'FFCD00' else GRIS_TITRE

doc.add_paragraph()
p_footer = add_para(doc,
    'Document préparé par l\'équipe projet PNPI — Version 2.0 — Février 2026\n'
    'Ministère de l\'Industrie et de la Transformation Locale — République Gabonaise',
    align=WD_ALIGN_PARAGRAPH.CENTER, size=8.5, color=RGBColor(0x80,0x80,0x80), before=10)

# ══════════════════════════════════════════════════════════════════════════════
#  SAUVEGARDE
# ══════════════════════════════════════════════════════════════════════════════
output_path = r'c:\Users\MBA NDONG J B\pnpi\docs\PNPI_Note_Conceptuelle_v2.docx'
doc.save(output_path)
print(f'Document généré : {output_path}')
